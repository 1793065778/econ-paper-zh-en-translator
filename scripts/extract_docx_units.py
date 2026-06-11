#!/usr/bin/env python3
"""Extract ordered translation units from a DOCX manuscript."""

from __future__ import annotations

import argparse
import json
import re
import sys
from dataclasses import asdict, dataclass
from pathlib import Path


try:
    from docx import Document
except ImportError:  # pragma: no cover - exercised only when dependency is absent
    Document = None


SENTENCE_RE = re.compile(r"(?<=[。！？；!?;])\s*")
NUMBER_RE = re.compile(
    r"(?<![A-Za-z])(?:\d+(?:\.\d+)?%?|\d{4}年|p\s*[<=>]\s*0?\.\d+|R\^?2|N\s*=\s*\d+)",
    re.IGNORECASE,
)
VARIABLE_RE = re.compile(r"\b[A-Za-z][A-Za-z0-9_]{0,12}\b|[α-ωΑ-Ωβγδθλμρσφχψω]")
FIG_TABLE_RE = re.compile(r"\b(?:Table|Figure|Fig\.)\s*\d+|[表图]\s*\d+", re.IGNORECASE)
EQUATION_HINT_RE = re.compile(r"[=+\-*/∑Σ√≤≥≈×]|\\\(|\\\[")


@dataclass
class Unit:
    unit_id: str
    section_id: str
    section_title: str
    unit_type: str
    source_text: str
    sentences: list[str]
    protected_items: list[str]
    translate: bool
    needs_review: bool
    review_reason: str


def split_sentences(text: str) -> list[str]:
    parts = [part.strip() for part in SENTENCE_RE.split(text.strip()) if part.strip()]
    return parts or ([text.strip()] if text.strip() else [])


def unique(items: list[str]) -> list[str]:
    seen: set[str] = set()
    result: list[str] = []
    for item in items:
        if item and item not in seen:
            seen.add(item)
            result.append(item)
    return result


def protected_items(text: str) -> list[str]:
    items: list[str] = []
    items.extend(NUMBER_RE.findall(text))
    items.extend(VARIABLE_RE.findall(text))
    items.extend(FIG_TABLE_RE.findall(text))
    table_or_figure_numbers = set(re.findall(r"[表图]\s*(\d+)|(?:Table|Figure|Fig\.)\s*(\d+)", text, re.IGNORECASE))
    flattened_numbers = {number for pair in table_or_figure_numbers for number in pair if number}
    cleaned = []
    for item in items:
        item = item.strip()
        if item in flattened_numbers and re.search(rf"(?:[表图]\s*{re.escape(item)}|(?:Table|Figure|Fig\.)\s*{re.escape(item)})", text, re.IGNORECASE):
            continue
        cleaned.append(item)
    return unique(cleaned)


def is_heading(paragraph) -> bool:
    style_name = getattr(getattr(paragraph, "style", None), "name", "") or ""
    text = paragraph.text.strip()
    return style_name.lower().startswith("heading") or bool(re.match(r"^\d+(\.\d+)*\s+.{1,80}$", text))


def classify(text: str, paragraph, in_references: bool) -> str:
    stripped = text.strip()
    if not stripped:
        return "empty_or_formatting"
    if in_references:
        return "reference"
    if is_heading(paragraph):
        return "heading"
    if re.match(r"^(参考文献|References|Bibliography)\s*$", stripped, re.IGNORECASE):
        return "heading"
    if re.match(r"^([图表]\s*\d+|(?:Figure|Fig\.|Table)\s*\d+)", stripped, re.IGNORECASE):
        return "figure_or_table_note"
    if EQUATION_HINT_RE.search(stripped) and len(stripped) < 180:
        if re.search(r"[\u4e00-\u9fff]", stripped):
            return "equation_or_variable_explanation"
        return "formula_or_equation"
    return "body_paragraph"


def should_translate(unit_type: str) -> bool:
    return unit_type not in {"reference", "formula_or_equation", "empty_or_formatting"}


def review_reason(unit_type: str) -> str:
    if unit_type == "equation_or_variable_explanation":
        return "Contains equation symbols and Chinese explanatory text; confirm whether to translate as prose or preserve as formula."
    return ""


def count_table_content(document) -> tuple[int, int]:
    cell_count = 0
    paragraph_count = 0
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_count += 1
                paragraph_count += len([paragraph for paragraph in cell.paragraphs if paragraph.text.strip()])
    return cell_count, paragraph_count


def iter_paragraphs(document, include_table_cells: bool) -> list:
    paragraphs = list(document.paragraphs)
    if include_table_cells:
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.extend(cell.paragraphs)
    return paragraphs


def extract_units(source: Path, include_table_cells: bool = False) -> tuple[list[Unit], dict[str, object]]:
    if Document is None:
        raise RuntimeError("python-docx is required. Install it with: python -m pip install python-docx")

    document = Document(str(source))
    excluded_table_cells, excluded_table_paragraphs = (0, 0)
    if not include_table_cells:
        excluded_table_cells, excluded_table_paragraphs = count_table_content(document)

    units: list[Unit] = []
    section_id = "S01"
    section_title = ""
    section_index = 1
    paragraph_index = 0
    in_references = False

    for paragraph in iter_paragraphs(document, include_table_cells):
        text = paragraph.text.strip()
        if re.match(r"^(参考文献|References|Bibliography)\s*$", text, re.IGNORECASE):
            in_references = True

        unit_type = classify(text, paragraph, in_references)
        if unit_type == "heading" and text:
            section_index += 1 if paragraph_index else 0
            section_id = f"S{section_index:02d}"
            section_title = text
            paragraph_index = 0

        paragraph_index += 1
        unit_id = f"{section_id}-P{paragraph_index:03d}"
        units.append(
            Unit(
                unit_id=unit_id,
                section_id=section_id,
                section_title=section_title,
                unit_type=unit_type,
                source_text=text,
                sentences=split_sentences(text),
                protected_items=protected_items(text),
                translate=should_translate(unit_type),
                needs_review=unit_type == "equation_or_variable_explanation",
                review_reason=review_reason(unit_type),
            )
        )

    metadata = {
        "source": str(source),
        "include_table_cells": include_table_cells,
        "excluded_table_cells": excluded_table_cells,
        "excluded_table_paragraphs": excluded_table_paragraphs,
        "unit_count": len(units),
        "translatable_unit_count": len([unit for unit in units if unit.translate]),
        "review_unit_count": len([unit for unit in units if unit.needs_review]),
    }
    return units, metadata


def write_structure_map(units: list[Unit], metadata: dict[str, object], path: Path) -> None:
    lines = [
        "# Structure Map",
        "",
        "## Extraction Summary",
        "",
        f"- Source: `{metadata['source']}`",
        f"- Units: {metadata['unit_count']}",
        f"- Translatable units: {metadata['translatable_unit_count']}",
        f"- Units needing review: {metadata['review_unit_count']}",
        f"- Include table cells: {str(metadata['include_table_cells']).lower()}",
        f"- Excluded table cells: {metadata['excluded_table_cells']}",
        f"- Excluded non-empty table paragraphs: {metadata['excluded_table_paragraphs']}",
        "",
    ]
    current_section = None
    for unit in units:
        if unit.section_id != current_section:
            current_section = unit.section_id
            title = unit.section_title or "(untitled)"
            lines.extend([f"## {unit.section_id}: {title}", ""])
        lines.append(
            f"### {unit.unit_id} `{unit.unit_type}` translate={str(unit.translate).lower()} needs_review={str(unit.needs_review).lower()}"
        )
        if unit.review_reason:
            lines.append(f"- Review reason: {unit.review_reason}")
        if unit.protected_items:
            lines.append(f"- Protected items: {', '.join(unit.protected_items)}")
        lines.extend(["", "```text", unit.source_text, "```", ""])
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("source_docx", type=Path)
    parser.add_argument("--out-json", type=Path, default=Path("translation_units.json"))
    parser.add_argument("--out-map", type=Path, default=Path("structure_map.md"))
    parser.add_argument("--include-table-cells", action="store_true", help="Include Word table cell paragraphs. Default is to exclude them.")
    args = parser.parse_args(argv)

    try:
        units, metadata = extract_units(args.source_docx, include_table_cells=args.include_table_cells)
    except Exception as exc:
        print(f"extract_docx_units.py: {exc}", file=sys.stderr)
        return 2

    payload = {"metadata": metadata, "units": [asdict(unit) for unit in units]}
    json_text = json.dumps(payload, ensure_ascii=False, indent=2)
    json.loads(json_text)
    args.out_json.write_text(json_text, encoding="utf-8")
    write_structure_map(units, metadata, args.out_map)
    print(f"Wrote {len(units)} units to {args.out_json} and {args.out_map}")
    if metadata["excluded_table_paragraphs"]:
        print(f"Excluded {metadata['excluded_table_paragraphs']} non-empty table paragraphs by default.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
